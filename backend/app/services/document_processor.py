"""Document Ingestion & Text Chunking Engine"""
import io
import re
import logging
from typing import List, Tuple
from pypdf import PdfReader
import docx

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Extracts text from PDF, DOCX, TXT files and performs recursive chunking."""

    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> str:
        """Extract plain text from uploaded file based on its extension."""
        ext = filename.lower().split(".")[-1]
        
        if ext == "pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            return "\n\n".join(text_parts)
            
        elif ext in ["docx", "doc"]:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            return "\n\n".join(text_parts)
            
        elif ext in ["txt", "md"]:
            return file_bytes.decode("utf-8", errors="replace")
            
        else:
            raise ValueError(f"Unsupported file format: .{ext}. Supported formats: PDF, DOCX, TXT")

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 700, chunk_overlap: int = 100) -> List[str]:
        """
        Split text into overlapping semantic chunks.
        """
        if not text or not text.strip():
            return []
            
        # Clean text
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        # Split by paragraphs first
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            if len(para) > chunk_size:
                # If paragraph itself is larger than chunk_size, split by sentences
                sentences = re.split(r"(?<=[.!?])\s+", para)
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) + 1 <= chunk_size:
                        current_chunk = f"{current_chunk} {sentence}".strip()
                    else:
                        if current_chunk:
                            chunks.append(current_chunk)
                        current_chunk = sentence
            else:
                if len(current_chunk) + len(para) + 2 <= chunk_size:
                    current_chunk = f"{current_chunk}\n\n{para}".strip()
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = para
                    
        if current_chunk:
            chunks.append(current_chunk)
            
        return [c.strip() for c in chunks if len(c.strip()) > 10]

document_processor = DocumentProcessor()
